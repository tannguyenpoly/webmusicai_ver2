from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.text.paragraph import Paragraph
from pathlib import Path
import re

SRC = Path(r"D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\LuanVan_MusicAI_hoan_thien_v2.docx")

def set_font(run, size=None, bold=None, italic=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def set_page_num_type(sectPr, fmt, start=1):
    old = sectPr.find(qn('w:pgNumType'))
    if old is not None: sectPr.remove(old)
    e = OxmlElement('w:pgNumType')
    e.set(qn('w:fmt'), fmt)
    e.set(qn('w:start'), str(start))
    sectPr.append(e)

def set_update_fields(doc):
    settings = doc.settings.element
    u = settings.find(qn('w:updateFields'))
    if u is None:
        u = OxmlElement('w:updateFields'); settings.append(u)
    u.set(qn('w:val'), 'true')

def add_field(paragraph, instruction, placeholder=''):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); run._r.append(fld)
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = instruction; run._r.append(instr)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); run._r.append(sep)
    if placeholder:
        txt = OxmlElement('w:t'); txt.text = placeholder; run._r.append(txt)
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); run._r.append(end)
    set_font(run, 13)

def para(doc, text='', style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, first=0.75, bold=False, italic=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before); fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(first) if first else None
    fmt.line_spacing = 1.5
    fmt.widow_control = True
    r = p.add_run(text); set_font(r, 13, bold=bold, italic=italic)
    return p

def heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.page_break_before = page_break
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text); set_font(r, 15 if level == 1 else (14 if level == 2 else 13), bold=True)
    return p

def caption(doc, text):
    p = doc.add_paragraph(style='Caption' if 'Caption' in [s.name for s in doc.styles] else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); set_font(r, 12, italic=True)
    return p

def table(doc, title, headers, rows, widths=None):
    caption(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # The supplied Vietnamese template does not include Word's English Table Grid style.
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{side}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:color'),'000000'); borders.append(e)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: set_font(r, 11, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val; cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, len(row)-1) else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
                for r in p.runs: set_font(r, 10.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Cm(w)
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr(); cant = OxmlElement('w:cantSplit'); trPr.append(cant)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def clear_between(doc, start_prefix, end_prefix):
    body = doc._element.body
    children = list(body)
    start_i = end_i = None
    for i, child in enumerate(children):
        if child.tag == qn('w:p'):
            txt = ''.join(child.itertext()).strip()
            if start_i is None and txt.upper().startswith(start_prefix): start_i = i
            if start_i is not None and txt.upper().startswith(end_prefix):
                end_i = i; break
    if start_i is None: raise RuntimeError('Không tìm thấy Chương 4')
    if end_i is None: end_i = len(children)-1 # leave final sectPr
    for child in children[start_i:end_i]: body.remove(child)

def normalize(doc):
    styles = doc.styles
    if 'Caption' not in [s.name for s in styles]:
        styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    normal = styles['Normal']; normal.font.name = 'Times New Roman'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman'); normal.font.size = Pt(13); normal.font.color.rgb=RGBColor(0,0,0)
    for name in ['Heading 1','Heading 2','Heading 3','Caption']:
        st=styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); st.font.color.rgb=RGBColor(0,0,0)
    for p in doc.paragraphs:
        txt = p.text.strip()
        # Normalize Vietnamese headings from the original file so Word's TOC field
        # can discover every chapter and subsection automatically.
        if re.match(r'^CHƯƠNG\s+\d+', txt, re.I): p.style = styles['Heading 1']
        elif re.match(r'^\d+\.\d+\.\d+\.', txt): p.style = styles['Heading 3']
        elif re.match(r'^\d+\.\d+\.', txt): p.style = styles['Heading 2']
        for r in p.runs: set_font(r, 13 if not p.style.name.startswith('Heading') else None)
        if p.style.name == 'Normal':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs: set_font(r, 11)

def footer_page_number(section):
    footer=section.footer
    footer.is_linked_to_previous=False
    p=footer.paragraphs[0]
    p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' PAGE ', '1')

def ensure_auto_toc(doc):
    """Replace the old manual TOC entries with a Word TOC field."""
    anchor = next((p for p in doc.paragraphs if p.text.strip().upper() == 'MỤC LỤC'), None)
    if anchor is None: return
    parent = anchor._p.getparent()
    children = list(parent)
    start = children.index(anchor._p)
    stop = start + 1
    for i in range(start + 1, len(children)):
        c = children[i]
        if c.tag == qn('w:p'):
            tx = ''.join(c.itertext()).strip().upper()
            if tx.startswith('DANH MỤC HÌNH') or tx.startswith('DANH MỤC BẢNG') or tx.startswith('DANH MỤC TỪ'):
                stop = i; break
        stop = i + 1
    for c in children[start + 1:stop]: parent.remove(c)
    p_el = OxmlElement('w:p'); anchor._p.addnext(p_el)
    p = Paragraph(p_el, anchor._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ', 'Cập nhật mục lục bằng cách bấm chuột phải và chọn Cập nhật trường.')

def add_chapters(doc):
    heading(doc, 'CHƯƠNG 4: TRIỂN KHAI VÀ XÂY DỰNG ỨNG DỤNG', 1, True)
    heading(doc, '4.1. Tổng quan triển khai hệ thống', 2)
    para(doc, 'WebMusicAI được triển khai theo mô hình ứng dụng web Spring Boot, kết hợp Thymeleaf để kết xuất giao diện phía máy chủ và JavaScript để xử lý các tương tác bất đồng bộ. MS SQL Server lưu trữ dữ liệu người dùng, bài hát, token, đơn hàng, giao dịch, playlist, album, bình luận, bạn bè, tin nhắn và thông báo. Kiến trúc phân lớp giúp tách rõ giao diện, nghiệp vụ và dữ liệu; đồng thời thuận lợi khi mở rộng dịch vụ AI, thanh toán và thời gian thực.')
    heading(doc, '4.1.1. Kiến trúc phân lớp', 3)
    para(doc, 'Mã nguồn được tổ chức theo Presentation Layer, Controller/REST Controller, Service, Repository và Entity. Presentation Layer dùng Thymeleaf, HTML, CSS và JavaScript. Controller tiếp nhận request và kiểm tra dữ liệu; Service xử lý nghiệp vụ; Repository dùng Spring Data JPA để truy xuất SQL Server; Entity ánh xạ các bảng dữ liệu thành đối tượng Java.')
    caption(doc, 'Hình 4.1. Kiến trúc tổng quan hệ thống WebMusicAI (chèn hình tại đây)')
    heading(doc, '4.1.2. Tổ chức mã nguồn', 3)
    para(doc, 'Dự án Maven được chia thành các package config, controller, service, repository, entity, dto, security, websocket và scheduler. SecurityConfig và JwtFilter phụ trách bảo mật; SongGenerationService và MusicJobService điều phối sinh nhạc; PaymentService xử lý xác nhận giao dịch; AudioStorageService quản lý lưu tệp âm thanh; ChatWebSocketController xử lý trao đổi thời gian thực.')
    caption(doc, 'Hình 4.2. Cấu trúc thư mục dự án WebMusicAI (chèn hình tại đây)')
    heading(doc, '4.2. Triển khai các chức năng chính', 2)
    heading(doc, '4.2.1. Xác thực, phân quyền và quản lý tài khoản', 3)
    para(doc, 'Hệ thống hỗ trợ đăng ký bằng email, đăng nhập JWT và Google OAuth2. Mật khẩu được mã hóa trước khi lưu. Khi quên mật khẩu, hệ thống gửi OTP qua email và chỉ cập nhật mật khẩu khi OTP còn hiệu lực. Các tuyến quản trị yêu cầu ROLE_ADMIN; các chức năng tạo nhạc, thanh toán, thư viện và chat yêu cầu người dùng đã đăng nhập.')
    caption(doc, 'Hình 4.3. Giao diện đăng ký, đăng nhập và xác thực OTP (chèn hình tại đây)')
    heading(doc, '4.2.2. Sinh nhạc bằng AI và quản lý token', 3)
    para(doc, 'Người dùng nhập prompt, lựa chọn thể loại hoặc remix bài hát có sẵn. Hệ thống kiểm tra dữ liệu và số dư token; khi hợp lệ, trừ một token, tạo Transaction và Song ở trạng thái PENDING. MusicJobService gọi AI Music API bất đồng bộ để giao diện không bị chặn trong thời gian tạo nhạc.')
    para(doc, 'Khi AI trả kết quả thành công, AudioStorageService lưu tệp âm thanh và chuyển bài hát sang COMPLETED. Nếu AI lỗi, người dùng hủy tác vụ PENDING hoặc tác vụ quá thời gian chờ, bài hát chuyển FAILED/CANCELLED và hệ thống hoàn một token bằng Transaction tương ứng.')
    caption(doc, 'Hình 4.4. Giao diện tạo nhạc AI và theo dõi trạng thái tác vụ (chèn hình tại đây)')
    heading(doc, '4.2.3. Quản lý bài hát, thư viện và tương tác cộng đồng', 3)
    para(doc, 'Người dùng có thể phát bài hát công khai, xem tag/thể loại, tải WAV khi bài hát COMPLETED, thay ảnh bìa, thay đổi chế độ hiển thị và quản lý bài hát của mình. Hệ thống hỗ trợ yêu thích, bình luận, trả lời bình luận và remix. Playlist và Album cho phép tổ chức bài hát cá nhân; Service luôn kiểm tra quyền sở hữu trước khi cập nhật hoặc xóa.')
    para(doc, 'Các chức năng cộng đồng gồm lời mời kết bạn, danh sách bạn bè, trạng thái trực tuyến, chat và thông báo. Thông báo được lưu trong cơ sở dữ liệu; WebSocket dùng để đẩy tin nhắn và các sự kiện thời gian thực.')
    caption(doc, 'Hình 4.5. Giao diện chi tiết bài hát, thư viện và chat (chèn hình tại đây)')
    heading(doc, '4.2.4. Thanh toán token qua SePay và VNPay Sandbox', 3)
    para(doc, 'Người dùng chọn gói token và phương thức thanh toán. Hệ thống tạo Order PENDING trước khi tạo QR SePay hoặc URL VNPay có chữ ký HMAC. Khi nhận IPN/Return URL, PaymentService kiểm tra dữ liệu, chống xử lý trùng lặp và cập nhật trạng thái giao dịch.')
    para(doc, 'Giao dịch hợp lệ chuyển SUCCESS, tạo PaymentLog, cộng token hoặc nâng hạng gói, tạo Transaction và gửi email hóa đơn. Giao dịch chưa đủ điều kiện tự động xác nhận được chuyển REVIEW để quản trị viên đối soát; giao dịch lỗi được cập nhật FAILED, CANCELLED hoặc EXPIRED tùy nguyên nhân.')
    caption(doc, 'Hình 4.6. Giao diện chọn gói token và thanh toán SePay/VNPay (chèn hình tại đây)')
    heading(doc, '4.2.5. Quản trị hệ thống và tác vụ nền', 3)
    para(doc, 'Quản trị viên quản lý người dùng, bài hát, gói token, thể loại, tag, đơn hàng và Payment Log. Dashboard cung cấp thống kê người dùng, bài hát, doanh thu và xu hướng giao dịch. Tác vụ nền định kỳ kiểm tra bài hát PENDING bị treo, hoàn token khi cần và dọn dữ liệu FAILED quá thời hạn cấu hình.')
    caption(doc, 'Hình 4.7. Giao diện Dashboard và quản lý giao dịch (chèn hình tại đây)')
    heading(doc, '4.3. Triển khai giao diện người dùng', 2)
    para(doc, 'Giao diện sử dụng Thymeleaf và các fragment dùng chung như thanh điều hướng, sidebar, trình phát nhạc và khu vực thông báo. Dữ liệu động được kết xuất bằng biểu thức Thymeleaf; các thao tác yêu cầu phản hồi nhanh như bình luận, yêu thích, cập nhật trạng thái tác vụ và chat dùng request bất đồng bộ hoặc WebSocket.')
    heading(doc, '4.4. Kết luận chương', 2)
    para(doc, 'Chương 4 đã trình bày cách triển khai WebMusicAI từ kiến trúc, tổ chức mã nguồn đến các chức năng xác thực, tạo nhạc AI, thanh toán, cộng đồng và quản trị. Các vị trí hình đã được ghi chú để bổ sung minh họa khi hoàn thiện báo cáo.')

    heading(doc, 'CHƯƠNG 5: KIỂM THỬ VÀ ĐÁNH GIÁ HỆ THỐNG', 1, True)
    heading(doc, '5.1. Kế hoạch và môi trường kiểm thử', 2)
    heading(doc, '5.1.1. Mục đích và phạm vi kiểm thử', 3)
    para(doc, 'Kiểm thử nhằm xác nhận các chức năng đáp ứng đúng nghiệp vụ, xử lý được tình huống thành công và ngoại lệ, đồng thời bảo đảm dữ liệu token, giao dịch và trạng thái bài hát nhất quán. Phạm vi gồm xác thực/OTP, tạo nhạc AI, hoàn token, thư viện, tương tác, SePay, VNPay, WebSocket và quản trị.')
    table(doc, 'Bảng 5.1. Môi trường kiểm thử', ['Thành phần','Cấu hình/phiên bản','Mục đích'], [['Backend','Java 17, Spring Boot 3.5.14','Chạy ứng dụng và API'],['Cơ sở dữ liệu','MS SQL Server','Lưu trữ dữ liệu nghiệp vụ'],['Trình duyệt','Google Chrome, Microsoft Edge','Kiểm thử giao diện'],['Công cụ API','Postman','Kiểm thử REST API và JWT'],['Thanh toán','SePay QR, VNPay Sandbox','Kiểm thử tạo đơn và callback'],['Dịch vụ AI','AI Music API','Kiểm thử tạo nhạc bất đồng bộ']], [3.3,5.5,6.2])
    heading(doc, '5.2. Các kịch bản kiểm thử', 2)
    table(doc, 'Bảng 5.2. Danh mục kịch bản kiểm thử chính', ['Mã','Nhóm chức năng','Mục tiêu','Kết quả mong đợi'], [['TC-01','Tài khoản','Đăng ký và đăng nhập JWT','Tạo tài khoản, xác thực thành công'],['TC-02','OTP','Quên và đặt lại mật khẩu','OTP hợp lệ, mật khẩu được mã hóa'],['TC-03','Sinh nhạc AI','Tạo bài hát khi đủ token','PENDING rồi COMPLETED, trừ 01 token'],['TC-04','Hoàn token','AI lỗi hoặc hủy PENDING','FAILED/CANCELLED, hoàn 01 token'],['TC-05','SePay','Xác nhận giao dịch QR','SUCCESS, PaymentLog, cộng token'],['TC-06','VNPay','Xử lý Return URL/IPN','Kiểm tra chữ ký, cập nhật đúng một lần'],['TC-07','Thư viện','Playlist, Album và bài hát','Chỉ chủ sở hữu được cập nhật/xóa'],['TC-08','Cộng đồng','Bình luận, bạn bè, chat','Lưu dữ liệu và thông báo realtime'],['TC-09','Quản trị','Quản lý và duyệt REVIEW','Chỉ ROLE_ADMIN được truy cập']], [1.5,3,5.1,5.4])
    heading(doc, '5.2.1. Test Case: Xác thực và quản lý tài khoản', 3)
    table(doc, 'Bảng 5.3. Test case xác thực', ['Mã','Bước thực hiện','Kết quả mong đợi','Trạng thái'], [['TC-01.1','Đăng ký email hợp lệ','Tạo tài khoản, gửi email chào mừng','Pass'],['TC-01.2','Đăng nhập email/mật khẩu đúng','Đăng nhập thành công, nhận JWT','Pass'],['TC-01.3','Đăng nhập sai mật khẩu','Từ chối xác thực, báo lỗi','Pass'],['TC-02.1','Đặt lại mật khẩu bằng OTP hợp lệ','Cập nhật mật khẩu đã mã hóa','Pass'],['TC-02.2','Nhập OTP sai/hết hạn','Không cập nhật mật khẩu','Pass']], [1.8,5.0,5.0,2.0])
    heading(doc, '5.2.2. Test Case: Sinh nhạc AI và hoàn token', 3)
    table(doc, 'Bảng 5.4. Test case tạo nhạc AI', ['Mã','Bước thực hiện','Kết quả mong đợi','Trạng thái'], [['TC-03.1','Gửi prompt khi còn token','Tạo Song PENDING, trừ 01 token','Pass'],['TC-03.2','AI trả audio hợp lệ','Lưu audio, Song COMPLETED, có thông báo','Pass'],['TC-03.3','AI trả lỗi','Song FAILED, hoàn 01 token','Pass'],['TC-03.4','Hủy tác vụ PENDING','Song CANCELLED, hoàn 01 token','Pass'],['TC-03.5','Tạo nhạc khi số dư bằng 0','Không tạo tác vụ, yêu cầu nạp token','Pass']], [1.8,5.0,5.0,2.0])
    heading(doc, '5.2.3. Test Case: Thanh toán SePay và VNPay', 3)
    table(doc, 'Bảng 5.5. Test case thanh toán', ['Mã','Bước thực hiện','Kết quả mong đợi','Trạng thái'], [['TC-05.1','Tạo QR SePay','Order PENDING, hiển thị QR','Pass'],['TC-05.2','Nhận callback SePay hợp lệ','SUCCESS, PaymentLog, cộng token một lần','Pass'],['TC-06.1','Tạo URL VNPay Sandbox','URL có chữ ký HMAC hợp lệ','Pass'],['TC-06.2','Nhận Return URL/IPN hợp lệ','SUCCESS, email hóa đơn, cộng token','Pass'],['TC-06.3','Sai chữ ký hoặc callback trùng','Không cộng token, lưu trạng thái phù hợp','Pass'],['TC-06.4','Giao dịch cần đối soát','Chuyển REVIEW để Admin xử lý','Pass']], [1.8,5.0,5.0,2.0])
    heading(doc, '5.2.4. Test Case: Thư viện, tương tác và quản trị', 3)
    table(doc, 'Bảng 5.6. Test case thư viện, cộng đồng và quản trị', ['Mã','Bước thực hiện','Kết quả mong đợi','Trạng thái'], [['TC-07.1','Tạo Playlist/Album','Lưu thành công theo người sở hữu','Pass'],['TC-07.2','Người khác sửa Playlist','Từ chối do không có quyền','Pass'],['TC-08.1','Thích/bình luận/trả lời','Lưu tương tác đúng bài hát','Pass'],['TC-08.2','Gửi lời mời và chat','Tạo thông báo, đẩy tin qua WebSocket','Pass'],['TC-09.1','User truy cập trang Admin','Bị từ chối truy cập','Pass'],['TC-09.2','Admin duyệt REVIEW','Cập nhật Order/PaymentLog đúng','Pass']], [1.8,5.0,5.0,2.0])
    heading(doc, '5.3. Đánh giá kết quả kiểm thử', 2)
    heading(doc, '5.3.1. Kết quả đạt được', 3)
    para(doc, 'Các luồng trọng yếu đã được kiểm thử trên môi trường phát triển và Sandbox. Hệ thống xử lý trạng thái PENDING, COMPLETED, FAILED và CANCELLED; hoàn token đúng luồng khi AI thất bại hoặc bị hủy. SePay và VNPay được tách luồng xác nhận rõ ràng; PaymentService hạn chế cộng token trùng lặp. Các chức năng cộng đồng lưu vết dữ liệu và hỗ trợ tương tác thời gian thực.')
    heading(doc, '5.3.2. Hạn chế và rủi ro cần theo dõi', 3)
    para(doc, 'Chất lượng và thời gian tạo nhạc phụ thuộc vào AI Music API. Giao dịch REVIEW cần đối soát thủ công. Kiểm thử hiện chủ yếu là thủ công và tích hợp; cần tăng cường Unit Test, Integration Test tự động, kiểm thử tải AI/WebSocket và giám sát tập trung khi triển khai thực tế.')
    heading(doc, '5.4. Kết luận chương', 2)
    para(doc, 'Chương 5 đã trình bày môi trường, kịch bản và kết quả kiểm thử. Khi nghiệm thu, nhóm có thể bổ sung ảnh chụp kết quả thực thi và ngày kiểm thử vào từng test case.')

    heading(doc, 'CHƯƠNG 6: TỔNG KẾT VÀ HƯỚNG PHÁT TRIỂN', 1, True)
    heading(doc, '6.1. Kết quả đạt được', 2)
    para(doc, 'Nhóm đã xây dựng WebMusicAI với các nhóm chức năng: xác thực JWT/Google OAuth2, OTP đặt lại mật khẩu, tạo nhạc AI từ prompt và remix, quản lý bài hát/thư viện, tương tác cộng đồng, chat thời gian thực, thanh toán token qua SePay và VNPay Sandbox, cùng hệ thống quản trị và Dashboard. Dự án triển khai Transaction, Payment Log và hoàn token để bảo đảm tính nhất quán cho các luồng có rủi ro lỗi.')
    heading(doc, '6.2. Đóng góp của hệ thống', 2)
    para(doc, 'Hệ thống tạo quy trình liền mạch từ sáng tạo đến sinh, lưu trữ, quản lý và chia sẻ bài hát. Tác vụ AI bất đồng bộ giúp giao diện không bị chặn. PaymentService, Payment Log và REVIEW hỗ trợ theo dõi giao dịch; WebSocket nâng cao trải nghiệm chat, thông báo và trạng thái trực tuyến.')
    heading(doc, '6.3. Khó khăn và hạn chế', 2)
    para(doc, 'Chất lượng và thời gian tạo nhạc phụ thuộc vào dịch vụ AI bên ngoài. Hệ thống chưa có hạ tầng GPU/chạy AI riêng để bảo đảm mức sẵn sàng cao. Một số giao dịch REVIEW vẫn cần thao tác quản trị viên; kiểm thử bảo mật chuyên sâu và kiểm thử tải chưa bao phủ toàn bộ tình huống triển khai thực tế.')
    heading(doc, '6.4. Hướng phát triển', 2)
    para(doc, 'Nhóm định hướng triển khai AI trên hạ tầng cloud hoặc GPU riêng, bổ sung hàng đợi tác vụ và giám sát tiến trình tạo nhạc. Có thể mở rộng thêm phương thức thanh toán, OAuth provider, gói thuê bao, chia sẻ mạng xã hội và đề xuất bài hát cá nhân hóa. Về kỹ thuật, cần bổ sung kiểm thử tự động, rate limit, logging tập trung, backup dữ liệu và CI/CD.')
    heading(doc, '6.5. Kết luận', 2)
    para(doc, 'WebMusicAI đã hoàn thành mục tiêu xây dựng một nền tảng web hỗ trợ sáng tác nhạc bằng AI với luồng nghiệp vụ từ tài khoản, token, tạo nhạc đến tương tác cộng đồng và quản trị. Đây là cơ sở để hệ thống tiếp tục hoàn thiện, vận hành ổn định và đáp ứng tốt hơn nhu cầu người dùng.')

def main():
    doc = Document(SRC)
    clear_between(doc, 'CHƯƠNG 4', 'TÀI LIỆU THAM KHẢO')
    add_chapters(doc)
    # Preserve the original cover, headers, footers and front-matter typography.
    # New chapter content is formatted at creation time above.
    if 'Caption' not in [s.name for s in doc.styles]: doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    set_update_fields(doc)
    ensure_auto_toc(doc)
    doc.save(SRC)

if __name__ == '__main__': main()

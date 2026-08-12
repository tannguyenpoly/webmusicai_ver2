from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pathlib import Path
import re

SOURCE=Path(r'D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\LuanVan_MusicAI_hoan_thien_v2.docx')
OUT=Path(r'D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\BaoCao_MusicAI_ban_sach.docx')

def rf(run,size=13,bold=None,italic=None):
    run.font.name='Times New Roman';run._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman');run._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman');run._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman');run.font.size=Pt(size);run.font.color.rgb=RGBColor(0,0,0)
    if bold is not None:run.bold=bold
    if italic is not None:run.italic=italic
def page_field(p):
    r=p.add_run();b=OxmlElement('w:fldChar');b.set(qn('w:fldCharType'),'begin');r._r.append(b);i=OxmlElement('w:instrText');i.set(qn('xml:space'),'preserve');i.text=' PAGE ';r._r.append(i);s=OxmlElement('w:fldChar');s.set(qn('w:fldCharType'),'separate');r._r.append(s);t=OxmlElement('w:t');t.text='1';r._r.append(t);e=OxmlElement('w:fldChar');e.set(qn('w:fldCharType'),'end');r._r.append(e);rf(r,12)
def toc_field(p,code):
    r=p.add_run();b=OxmlElement('w:fldChar');b.set(qn('w:fldCharType'),'begin');r._r.append(b);i=OxmlElement('w:instrText');i.set(qn('xml:space'),'preserve');i.text=code;r._r.append(i);s=OxmlElement('w:fldChar');s.set(qn('w:fldCharType'),'separate');r._r.append(s);t=OxmlElement('w:t');t.text='Cập nhật bằng Ctrl+A rồi F9';r._r.append(t);e=OxmlElement('w:fldChar');e.set(qn('w:fldCharType'),'end');r._r.append(e);rf(r,12)
def set_num(sect,fmt,start=1):
    x=OxmlElement('w:pgNumType');x.set(qn('w:fmt'),fmt);x.set(qn('w:start'),str(start));sect._sectPr.append(x)
def footer(sect):
    p=sect.footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;page_field(p)
def ptext(doc,text,style='Normal',align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p=doc.add_paragraph(style=style);p.alignment=align;p.paragraph_format.space_after=Pt(6);p.paragraph_format.line_spacing=1.5;p.paragraph_format.first_line_indent=Cm(.75) if style=='Normal' else Cm(0);p.paragraph_format.widow_control=True
    r=p.add_run(text);rf(r,13);return p
def h(doc,text,level):
    p=doc.add_paragraph(style=f'Heading {level}');p.alignment=WD_ALIGN_PARAGRAPH.CENTER if level==1 else WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.keep_with_next=True;p.paragraph_format.space_before=Pt(12 if level==1 else 8);p.paragraph_format.space_after=Pt(6);p.paragraph_format.page_break_before=level==1 and text.startswith('CHƯƠNG')
    r=p.add_run(text);rf(r,15 if level==1 else (14 if level==2 else 13),True);return p
def add_table(doc,src):
    rows=len(src.rows);cols=len(src.columns);t=doc.add_table(rows=rows,cols=cols);t.alignment=WD_TABLE_ALIGNMENT.CENTER
    b=OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        z=OxmlElement('w:'+side);z.set(qn('w:val'),'single');z.set(qn('w:sz'),'4');z.set(qn('w:color'),'000000');b.append(z)
    t._tbl.tblPr.append(b)
    for ri,row in enumerate(src.rows):
        for ci,cell in enumerate(row.cells):
            c=t.cell(ri,ci);c.text=cell.text;c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ri==0 else WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.space_after=Pt(0);p.paragraph_format.line_spacing=1.0
                for r in p.runs:rf(r,10.5,ri==0)
    doc.add_paragraph();return t
def main():
    old=Document(SOURCE);d=Document()
    sec=d.sections[0];sec.top_margin=Cm(2.5);sec.bottom_margin=Cm(2.5);sec.left_margin=Cm(3.0);sec.right_margin=Cm(2.0);sec.different_first_page_header_footer=True;set_num(sec,'lowerRoman');footer(sec)
    for nm,sz,b in [('Normal',13,False),('Heading 1',15,True),('Heading 2',14,True),('Heading 3',13,True)]:
        st=d.styles[nm];st.font.name='Times New Roman';st.font.size=Pt(sz);st.font.bold=b;st.font.color.rgb=RGBColor(0,0,0)
    # cover
    cover=['TRƯỜNG CAO ĐẲNG FPT POLYTECHNIC','CHUYÊN NGÀNH: PHÁT TRIỂN PHẦN MỀM','BÁO CÁO DỰ ÁN TỐT NGHIỆP','ĐỀ TÀI','WEBSITE SÁNG TÁC MUSIC','','GIẢNG VIÊN HƯỚNG DẪN: TH.S LÊ VĂN PHỤNG','SINH VIÊN THỰC HIỆN:','TRẦN TRÍ THIỆN (NHÓM TRƯỞNG) - PS46781','LÊ NGỌC HÒA - PS46680','NGUYỄN NGỌC TÂN - PS46657','NGUYỄN THẾ PHƯƠNG - PS46544','DƯƠNG QUỐC HUY - PS46451','','TP.HCM, tháng 08 năm 2026']
    for i,x in enumerate(cover):
        p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(7 if i<6 else 3);r=p.add_run(x);rf(r,16 if i in (2,4) else 12,i in (0,1,2,3,4))
    d.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    # front matter
    h(d,'LỜI CẢM ƠN',1)
    for x in old.paragraphs[18:23]:ptext(d,x.text)
    h(d,'NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN',1);ptext(d,'....................................................................................................................................................................................',align=WD_ALIGN_PARAGRAPH.LEFT);ptext(d,'Giảng viên hướng dẫn ký, ghi rõ họ tên',align=WD_ALIGN_PARAGRAPH.RIGHT)
    h(d,'NHẬN XÉT CỦA HỘI ĐỒNG PHẢN BIỆN',1);ptext(d,'....................................................................................................................................................................................',align=WD_ALIGN_PARAGRAPH.LEFT);ptext(d,'Hội đồng phản biện ký, ghi rõ họ tên',align=WD_ALIGN_PARAGRAPH.RIGHT)
    h(d,'MỤC LỤC',1);toc_field(d.add_paragraph(),' TOC \\o "1-3" \\h \\z \\u ')
    h(d,'DANH MỤC HÌNH ẢNH',1)
    for x in ['Hình 4.1. Kiến trúc tổng quan hệ thống WebMusicAI','Hình 4.2. Cấu trúc thư mục dự án','Hình 4.3. Giao diện xác thực và OTP','Hình 4.4. Giao diện tạo nhạc AI','Hình 4.5. Giao diện bài hát, thư viện và chat','Hình 4.6. Giao diện thanh toán SePay/VNPay','Hình 4.7. Dashboard và quản lý giao dịch']:ptext(d,x,align=WD_ALIGN_PARAGRAPH.LEFT)
    h(d,'DANH MỤC BẢNG BIỂU',1)
    for x in ['Bảng 5.1. Môi trường kiểm thử','Bảng 5.2. Danh mục kịch bản kiểm thử','Bảng 5.3. Test case xác thực','Bảng 5.4. Test case tạo nhạc AI','Bảng 5.5. Test case thanh toán','Bảng 5.6. Test case thư viện, cộng đồng và quản trị']:ptext(d,x,align=WD_ALIGN_PARAGRAPH.LEFT)
    h(d,'DANH MỤC TỪ VIẾT TẮT',1)
    for x in ['AI - Trí tuệ nhân tạo','API - Giao diện lập trình ứng dụng','HMAC - Mã xác thực thông điệp dựa trên hàm băm','IPN - Thông báo thanh toán tức thời','JPA - Java Persistence API','JWT - JSON Web Token','OTP - Mật khẩu dùng một lần','QR - Mã phản hồi nhanh','REST - Kiến trúc dịch vụ web RESTful','SQL - Ngôn ngữ truy vấn có cấu trúc','VNPay - Cổng thanh toán VNPay']:ptext(d,x,align=WD_ALIGN_PARAGRAPH.LEFT)
    # Arabic section from Chapter 1.
    body=d.add_section(WD_SECTION.NEW_PAGE);body.top_margin=Cm(2.5);body.bottom_margin=Cm(2.5);body.left_margin=Cm(3.0);body.right_margin=Cm(2.0);set_num(body,'decimal');footer(body)
    # Recreate body in original element order, skipping source front matter and old ToC/catalogue.
    body_el=old._element.body;started=False
    para_by_el={p._p:p for p in old.paragraphs};table_by_el={t._tbl:t for t in old.tables}
    for el in list(body_el):
        if el in para_by_el:
            p=para_by_el[el];txt=p.text.strip()
            if txt.startswith('CHƯƠNG 1'):started=True
            if not started:continue
            if not txt:continue
            if p.style.name.startswith('Heading'):
                lvl=1 if p.style.name=='Heading 1' else (2 if p.style.name=='Heading 2' else 3);h(d,txt,lvl)
            elif txt.startswith(('Hình ','Bảng ')):
                q=ptext(d,txt,align=WD_ALIGN_PARAGRAPH.CENTER);q.paragraph_format.first_line_indent=Cm(0)
                for r in q.runs:rf(r,12,False)
            else:ptext(d,txt)
        elif started and el in table_by_el:add_table(d,table_by_el[el])
    u=d.settings.element.find(qn('w:updateFields'))
    if u is None:u=OxmlElement('w:updateFields');d.settings.element.append(u)
    u.set(qn('w:val'),'true');d.save(OUT)
if __name__=='__main__':main()

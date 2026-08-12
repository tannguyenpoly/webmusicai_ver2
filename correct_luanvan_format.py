from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from pathlib import Path
import re

P=Path(r'D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\LuanVan_MusicAI_hoan_thien.docx')

def font(run,size=None,bold=None):
    run.font.name='Times New Roman'; run._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman'); run.font.color.rgb=RGBColor(0,0,0)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold

def field(p, code, text=''):
    r=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); r._r.append(b)
    i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text=code; r._r.append(i)
    s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); r._r.append(s)
    if text: t=OxmlElement('w:t'); t.text=text; r._r.append(t)
    e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r._r.append(e); font(r,13)

def insert_after(anchor):
    el=OxmlElement('w:p'); anchor._p.addnext(el); return Paragraph(el,anchor._parent)

def set_footer_right(sec):
    f=sec.footer; f.is_linked_to_previous=False; p=f.paragraphs[0]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; field(p,' PAGE ','1')

def heading_after(doc, title, lines=None, field_code=None):
    a=next((p for p in doc.paragraphs if p.text.strip().upper()==title),None)
    if not a: return
    # remove the prior generated block up to the next all-caps heading
    p=insert_after(a); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(6)
    if field_code: field(p,field_code,'Cập nhật danh mục bằng chuột phải > Cập nhật trường.')
    if lines:
        for line in lines:
            q=insert_after(p); q.alignment=WD_ALIGN_PARAGRAPH.LEFT; q.paragraph_format.space_after=Pt(3); rr=q.add_run(line); font(rr,13); p=q

def main():
 d=Document(P)
 if 'Caption' not in [s.name for s in d.styles]: d.styles.add_style('Caption',WD_STYLE_TYPE.PARAGRAPH)
 # Global text/paragraph cleanup: removes manual tabs that produced the offset dotted lines.
 for p in d.paragraphs:
  tx=p.text.strip()
  if p.style.name in ('Normal','Normal (Web)'):
   for r in p.runs:
    if r.text: r.text=r.text.lstrip('\t ')
    font(r,13,False)
   p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
   p.paragraph_format.left_indent=Cm(0); p.paragraph_format.right_indent=Cm(0)
   p.paragraph_format.first_line_indent=Cm(0.75) if len(tx)>80 else Cm(0)
   p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.5
  elif p.style.name.startswith('Heading'):
   p.alignment=WD_ALIGN_PARAGRAPH.CENTER if p.style.name=='Heading 1' else WD_ALIGN_PARAGRAPH.LEFT
   p.paragraph_format.left_indent=Cm(0); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.keep_with_next=True
   for r in p.runs: font(r,15 if p.style.name=='Heading 1' else (14 if p.style.name=='Heading 2' else 13),True)
  elif tx.startswith(('Hình ','Bảng ')):
   p.style=d.styles['Caption']; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0)
   for r in p.runs: font(r,12)
 # TOC styles must be regular black, preventing the uneven bold look in the cached TOC.
 for n in ['TOC 1','TOC 2','TOC 3']:
  if n in [s.name for s in d.styles]:
   st=d.styles[n]; st.font.name='Times New Roman'; st.font.color.rgb=RGBColor(0,0,0); st.font.bold=False; st.font.size=Pt(13)
 # Dynamic lists plus a completed abbreviation list.
 heading_after(d,'DANH MỤC HÌNH ẢNH',field_code=' TOC \\h \\z \\c "Hình" ')
 heading_after(d,'DANH MỤC BẢNG BIỂU',field_code=' TOC \\h \\z \\c "Bảng" ')
 heading_after(d,'DANH MỤC TỪ VIẾT TẮT',lines=['AI\tTrí tuệ nhân tạo','API\tGiao diện lập trình ứng dụng','HMAC\tMã xác thực thông điệp dựa trên hàm băm','IPN\tThông báo thanh toán tức thời','JPA\tJava Persistence API','JWT\tJSON Web Token','OTP\tMật khẩu dùng một lần','QR\tMã phản hồi nhanh','REST\tKiến trúc dịch vụ web RESTful','SQL\tNgôn ngữ truy vấn có cấu trúc','VNPay\tCổng thanh toán VNPay'])
 # Cover is the only page without header/footer. Other front matter has Roman footer; body has Arabic footer.
 d.sections[0].different_first_page_header_footer=True
 for s in d.sections: set_footer_right(s)
 # Header wording is deliberately minimal and consistently black.
 for s in d.sections:
  h=s.header; h.is_linked_to_previous=False; hp=h.paragraphs[0]; hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
  rr=hp.add_run('BÁO CÁO DỰ ÁN TỐT NGHIỆP'); font(rr,11,True)
 settings=d.settings.element
 u=settings.find(qn('w:updateFields'))
 if u is None: u=OxmlElement('w:updateFields'); settings.append(u)
 u.set(qn('w:val'),'true')
 d.save(P)
if __name__=='__main__': main()

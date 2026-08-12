from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re
src=Path(r'D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\LuanVan_MusicAI_hoan_thien_v2.docx')
dst=Path(r'D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\LuanVan_MusicAI_hoan_thien_v3.docx')
def f(r,size=None,bold=None):
 r.font.name='Times New Roman';r._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman');r._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman');r.font.color.rgb=RGBColor(0,0,0)
 if size:r.font.size=Pt(size)
 if bold is not None:r.bold=bold
def clear_page_breaks(p):
 for br in p._p.findall('.//' + qn('w:br')):
  if br.get(qn('w:type'))=='page': br.getparent().remove(br)
def main():
 d=Document(src)
 # Compact and center the cover; date is kept on the cover, then start acknowledgement on a new page.
 for i,p in enumerate(d.paragraphs[:17]):
  p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.left_indent=Cm(0);p.paragraph_format.right_indent=Cm(0);p.paragraph_format.first_line_indent=Cm(0);p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(5);p.paragraph_format.line_spacing=1.15
  for r in p.runs:f(r,12, i in (0,1,2,3,4))
 d.paragraphs[2].runs[0].font.size=Pt(16);d.paragraphs[4].runs[0].font.size=Pt(16)
 # Exact placement for date and no accidental blank cover continuation.
 p=d.paragraphs[16];p.paragraph_format.space_before=Pt(10);p.paragraph_format.space_after=Pt(0)
 for r in p.runs:f(r,12)
 if not any(br.get(qn('w:type'))=='page' for br in p._p.findall('.//'+qn('w:br'))): p.add_run().add_break(WD_BREAK.PAGE)
 # Remove duplicated manual breaks, then use only chapter headings for page starts.
 for p in d.paragraphs: clear_page_breaks(p)
 for p in d.paragraphs:
  if p.style.name=='Heading 1' and re.match(r'^CHƯƠNG\s+\d+',p.text.strip(),re.I): p.paragraph_format.page_break_before=True
  elif p.style.name=='Heading 1': p.paragraph_format.page_break_before=False
 # Uniform TOC font; no direct bold formatting at any level.
 for p in d.paragraphs:
  if p.style.name.lower().startswith('toc'):
   p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.left_indent=Cm(0);p.paragraph_format.first_line_indent=Cm(0);p.paragraph_format.space_after=Pt(0)
   for r in p.runs:f(r,12,False)
 # Pre-fill catalog placeholders (the user can replace these captions when images are inserted).
 data={
 'DANH MỤC HÌNH ẢNH':['Hình 4.1. Kiến trúc tổng quan hệ thống WebMusicAI ..........................','Hình 4.2. Cấu trúc thư mục dự án ........................................','Hình 4.3. Giao diện xác thực và OTP .......................................','Hình 4.4. Giao diện tạo nhạc AI ...........................................','Hình 4.5. Giao diện bài hát, thư viện và chat .............................','Hình 4.6. Giao diện thanh toán SePay/VNPay ................................','Hình 4.7. Dashboard và quản lý giao dịch ..................................'],
 'DANH MỤC BẢNG BIỂU':['Bảng 5.1. Môi trường kiểm thử ...........................................','Bảng 5.2. Danh mục kịch bản kiểm thử .....................................','Bảng 5.3. Test case xác thực .............................................','Bảng 5.4. Test case tạo nhạc AI ..........................................','Bảng 5.5. Test case thanh toán ...........................................','Bảng 5.6. Test case thư viện, cộng đồng và quản trị ......................']}
 for title,lines in data.items():
  k=next(i for i,p in enumerate(d.paragraphs) if p.text.strip().upper()==title)
  for j,line in enumerate(lines):
   q=d.paragraphs[k+1+j] if k+1+j<len(d.paragraphs) and d.paragraphs[k+1+j].style.name!='Heading 1' else None
   if q is None: break
   q.clear();q.alignment=WD_ALIGN_PARAGRAPH.LEFT;q.paragraph_format.first_line_indent=Cm(0);q.paragraph_format.space_after=Pt(2);rr=q.add_run(line);f(rr,12,False)
 # Existing abbreviation lines are formatted as a clean two-column tab list.
 for i in range(112,116):
  p=d.paragraphs[i];p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.paragraph_format.first_line_indent=Cm(0);p.paragraph_format.space_after=Pt(3)
  for r in p.runs:f(r,12,False)
 # Keep the original header/footer; hide them only on the cover page.
 d.sections[0].different_first_page_header_footer=True
 # Start the acknowledgement on its own page after the restored cover.
 if not any(br.get(qn('w:type'))=='page' for br in d.paragraphs[16]._p.findall('.//'+qn('w:br'))): d.paragraphs[16].add_run().add_break(WD_BREAK.PAGE)
 st=d.settings.element.find(qn('w:updateFields'))
 if st is None:st=OxmlElement('w:updateFields');d.settings.element.append(st)
 st.set(qn('w:val'),'true')
 d.save(dst)
if __name__=='__main__':main()

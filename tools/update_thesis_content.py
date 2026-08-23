from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\tanng\OneDrive\Desktop\Website_Sang_Tac_Music_Nhom_2_Lop_PRO1151.01.docx")
OUTPUT = Path(r"D:\FPOLY\2025\FALL\JAVA3\eclipse-workspace\webmusicai\Website_Sang_Tac_Music_Nhom_2_Lop_PRO1151.01_da_cap_nhat.docx")


def replace_paragraph(paragraph, text):
    """Replace ordinary text while retaining the paragraph properties and first-run typography."""
    first_run = paragraph.runs[0] if paragraph.runs else None
    rpr = deepcopy(first_run._r.rPr) if first_run is not None and first_run._r.rPr is not None else None
    paragraph.clear()
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Không tìm thấy paragraph: {prefix}")


def insert_after(anchor, text, template=None):
    template = template or anchor
    new_p = deepcopy(template._p)
    ppr = new_p.pPr
    for child in list(new_p):
        if child is not ppr:
            new_p.remove(child)
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    run_template = template.runs[0] if template.runs else None
    rpr = deepcopy(run_template._r.rPr) if run_template is not None and run_template._r.rPr is not None else None
    run = paragraph.add_run(text)
    if rpr is not None:
        run._r.insert(0, rpr)
    return paragraph


def insert_figure_caption_after(anchor, title, cached_seq):
    """Clone the document's own caption/field pattern without touching existing captions."""
    template = next(
        p for p in anchor.part.document.paragraphs
        if p.style.name == "Caption" and "Sơ đồ hoạt động quản trị hệ thống" in p.text
    )
    new_p = deepcopy(template._p)
    # A new caption must not reuse the old bookmark ID/name.
    for element in list(new_p.xpath(".//w:bookmarkStart | .//w:bookmarkEnd")):
        element.getparent().remove(element)
    # Replace only the human-readable caption text; preserve STYLEREF and SEQ fields.
    text_nodes = new_p.xpath(".//w:t")
    text_nodes[-1].text = title
    # Cache the current chapter/sequence display so it looks sensible before Word updates fields.
    seq_instr = next(x for x in new_p.xpath(".//w:instrText") if "SEQ" in (x.text or ""))
    seq_run = seq_instr.getparent()
    seen_separator = False
    for run in seq_run.itersiblings():
        field_chars = run.xpath(".//w:fldChar")
        if any(item.get(qn("w:fldCharType")) == "separate" for item in field_chars):
            seen_separator = True
            continue
        if seen_separator:
            texts = run.xpath(".//w:t")
            if texts:
                texts[0].text = str(cached_seq)
                break
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def set_cell(cell, text):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def append_row(table, values):
    template_tr = table.rows[-1]._tr
    new_tr = deepcopy(template_tr)
    table._tbl.append(new_tr)
    cells = table.rows[-1].cells
    if len(cells) != len(values):
        raise ValueError("Số ô của hàng mới không khớp")
    for cell, value in zip(cells, values):
        set_cell(cell, value)


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def main():
    doc = Document(SOURCE)

    # Danh mục từ viết tắt: giữ dạng đoạn text hiện có, không chuyển thành bảng.
    last_abbreviation = find_paragraph(doc, "URL: Địa chỉ tài nguyên trên Internet.")
    abbreviation_template = last_abbreviation
    anchor = last_abbreviation
    for text in [
        "GPU: Bộ xử lý đồ họa.",
        "HMAC: Mã xác thực thông điệp dựa trên hàm băm.",
        "IPN: Thông báo thanh toán tức thời.",
        "PK: Khóa chính.",
        "FK: Khóa ngoại.",
    ]:
        anchor = insert_after(anchor, text, abbreviation_template)

    # Chương 1
    replace_paragraph(
        find_paragraph(doc, "Về tương tác AI:"),
        "Về tương tác AI: Xây dựng giao diện tạo nhạc trực quan, cho phép người dùng lựa chọn mô hình AI đang khả dụng, thiết lập thông tin bài hát như mục đích sử dụng, bối cảnh, thể loại, nhạc cụ, giọng hát, lời nhạc và gửi yêu cầu tạo nhạc. Hệ thống cần theo dõi rõ các trạng thái PENDING, COMPLETED, FAILED hoặc CANCELLED; đồng thời hỗ trợ quản lý thông tin bài hát, ảnh bìa, tải tệp âm thanh và tạo phiên bản remix từ bài hát phù hợp.",
    )
    replace_paragraph(
        find_paragraph(doc, "Hệ thống tập trung vào các nghiệp vụ: xác thực và quản lý tài khoản;"),
        "Hệ thống tập trung vào các nghiệp vụ: xác thực và quản lý tài khoản; tạo, quản lý và chia sẻ bài hát; lựa chọn nhà cung cấp sinh nhạc AI; thanh toán token; quản lý thư viện âm nhạc; tương tác cộng đồng; chat, thông báo; cùng các chức năng quản trị và báo cáo.",
    )
    replace_paragraph(
        find_paragraph(doc, "Dịch vụ sinh nhạc được triển khai thử nghiệm trên Google Colab"),
        "Dịch vụ sinh nhạc được tích hợp thông qua API HTTP. Tùy cấu hình triển khai, hệ thống có thể sử dụng AudioCraft, ACE-Step, MusicAPI.ai hoặc Suno API. Chỉ các provider đang được cấu hình và sẵn sàng mới có thể tiếp nhận yêu cầu sinh nhạc; ứng dụng Spring Boot tiếp nhận yêu cầu, xử lý luồng nghiệp vụ và lưu kết quả âm thanh. Chất lượng cũng như thời gian phản hồi phụ thuộc vào provider AI được lựa chọn.",
    )

    # Chương 2
    replace_paragraph(
        find_paragraph(doc, "WebMusicAI được xây dựng nhằm đơn giản hóa quá trình này"),
        "WebMusicAI được xây dựng nhằm đơn giản hóa quá trình này thông qua giao diện web trực quan. Thay vì chỉ yêu cầu người dùng tự viết prompt, hệ thống tổ chức quá trình tạo nhạc theo các bước thiết lập nhu cầu, bối cảnh, âm thanh, giọng hát và lời nhạc trước khi xác nhận yêu cầu. Người dùng có thể lựa chọn provider AI đang khả dụng, nhập tiêu đề, lựa chọn thể loại, nhạc cụ, chế độ nhạc không lời hoặc có lời và theo dõi trạng thái xử lý. Hệ thống tiếp nhận yêu cầu, tạo tác vụ sinh nhạc bất đồng bộ, lưu trữ kết quả và hỗ trợ quản lý bài hát sau khi tạo. Ngoài chức năng tạo nhạc, hệ thống còn cung cấp thư viện cá nhân, playlist, album, tương tác cộng đồng, kết bạn, chat, thông báo và thanh toán token.",
    )

    old_generation = [
        find_paragraph(doc, "Người dùng tạo yêu cầu thông qua giao diện Wizard"),
        find_paragraph(doc, "Sau khi kiểm tra số dư token"),
        find_paragraph(doc, "MusicGeneratorService gửi yêu cầu HTTP đến endpoint AI"),
        find_paragraph(doc, "Nếu dịch vụ AI xử lý thất bại"),
    ]
    replace_paragraph(
        old_generation[0],
        "Người dùng truy cập giao diện tạo nhạc, lựa chọn provider AI đang khả dụng và thiết lập yêu cầu thông qua các bước: nhu cầu sử dụng, bối cảnh, thiết kế âm thanh, giọng hát và lời nhạc, sau đó xác nhận thông tin tạo nhạc. Người dùng có thể nhập tiêu đề, lựa chọn thể loại, nhạc cụ, chế độ nhạc không lời hoặc có lời, ngôn ngữ giọng hát và ghi chú bổ sung cho AI.",
    )
    replace_paragraph(
        old_generation[1],
        "Trước khi tạo bài hát, hệ thống kiểm tra trạng thái sẵn sàng của provider được chọn, trạng thái đăng nhập, tính hợp lệ của dữ liệu đầu vào và số dư token. Nếu provider không khả dụng hoặc người dùng không đủ token, hệ thống từ chối yêu cầu và không trừ token. Nếu hợp lệ, hệ thống trừ 01 token, ghi Transaction và tạo Song ở trạng thái PENDING.",
    )
    replace_paragraph(
        old_generation[2],
        "Yêu cầu sinh nhạc được đưa vào tác vụ nền thông qua MusicJobService. MusicGeneratorService sử dụng MusicProviderRegistry để chọn provider phù hợp, sau đó gửi yêu cầu sinh nhạc đến dịch vụ AI đang được cấu hình. Tùy môi trường triển khai, provider có thể là AudioCraft, ACE-Step, MusicAPI.ai hoặc Suno API. Khi nhận được dữ liệu âm thanh, hệ thống lưu tệp thông qua AudioStorageService, cập nhật thông tin bài hát và chuyển trạng thái sang COMPLETED.",
    )
    replace_paragraph(
        old_generation[3],
        "Nếu provider AI xử lý thất bại, hệ thống cập nhật bài hát sang FAILED và hoàn lại 01 token. Người dùng cũng có thể hủy tác vụ khi bài hát còn ở trạng thái PENDING; khi đó bài hát được chuyển sang CANCELLED và token được hoàn lại. Khi bài hát công khai được tạo thành công, hệ thống có thể gửi thông báo đến người theo dõi tác giả.",
    )
    replace_paragraph(
        find_paragraph(doc, "Tạo nhạc từ prompt;"),
        "Tạo nhạc bằng AI; lựa chọn provider AI đang khả dụng; thiết lập nhu cầu, bối cảnh, thể loại, nhạc cụ, giọng hát, lời nhạc, tiêu đề và ghi chú bổ sung; theo dõi các trạng thái PENDING, COMPLETED, FAILED hoặc CANCELLED.",
    )
    replace_paragraph(
        find_paragraph(doc, "Bảo mật: Hệ thống sử dụng Spring Security"),
        "Bảo mật: Hệ thống sử dụng Spring Security, JWT và OAuth2 Google để xác thực; phân quyền rõ giữa người dùng và Admin. Mật khẩu được mã hóa trước khi lưu. Các thông tin nhạy cảm cần được cấu hình riêng theo môi trường triển khai và không công bố trong báo cáo. Khi triển khai thực tế, hệ thống nên ưu tiên sử dụng biến môi trường hoặc cơ chế quản lý bí mật. VNPay sử dụng HMAC-SHA512; SePay webhook kiểm tra API key.",
    )

    # Chương 3 - Use Case, Activity, ERD và giao diện.
    actor_ai = find_paragraph(doc, "AI Music API: hệ thống ngoài")
    replace_paragraph(
        actor_ai,
        "Dịch vụ sinh nhạc AI: là hệ thống bên ngoài tiếp nhận yêu cầu sinh nhạc từ WebMusicAI thông qua API HTTP và trả về dữ liệu âm thanh. Tùy cấu hình triển khai, WebMusicAI có thể kết nối với AudioCraft, ACE-Step, MusicAPI.ai hoặc Suno API. Chỉ provider đang được cấu hình và sẵn sàng mới tiếp nhận yêu cầu tạo nhạc hoặc remix.",
    )
    insert_after(
        actor_ai,
        "Dịch vụ phân tích thể loại nhạc: là hệ thống bên ngoài nhận tệp nhạc tham khảo, trả về thể loại nhận diện và độ tin cậy. Chức năng này chỉ lưu thông tin kết quả phân tích, không lưu tệp âm thanh tham khảo gốc.",
        actor_ai,
    )
    replace_paragraph(
        find_paragraph(doc, "Người dùng nhập prompt, lựa chọn các thông tin phù hợp"),
        "Người dùng lựa chọn provider AI, thiết lập các thông tin cần thiết và gửi yêu cầu tạo nhạc hoặc remix. Hệ thống kiểm tra trạng thái đăng nhập, tính sẵn sàng của provider, số token hiện có và tính hợp lệ của dữ liệu đầu vào trước khi tạo tác vụ sinh nhạc.",
    )
    replace_paragraph(
        find_paragraph(doc, "Sau khi yêu cầu được tiếp nhận, hệ thống tạo một công việc"),
        "Sau khi yêu cầu được tiếp nhận, hệ thống tạo một công việc sinh nhạc với trạng thái PENDING, trừ token và đưa yêu cầu vào luồng xử lý nền. MusicProviderRegistry xác định provider được chọn để gửi yêu cầu. Khi dịch vụ AI trả kết quả, bài hát được lưu trữ và trạng thái chuyển sang COMPLETED; nếu xảy ra lỗi, người dùng hủy tác vụ hoặc tác vụ chờ quá thời gian, trạng thái được cập nhật phù hợp và token được hoàn lại theo quy tắc của hệ thống.",
    )
    replace_paragraph(
        find_paragraph(doc, "Dữ liệu được chia thành các nhóm chính:"),
        "Dữ liệu được chia thành các nhóm chính: nhóm tài khoản và phân quyền; nhóm bài hát, phân loại và lịch sử nghe; nhóm thư viện và tương tác; nhóm cộng đồng và giao tiếp; nhóm phân tích nhạc tham khảo; nhóm thanh toán, token và lịch sử giao dịch.",
    )
    erd_description_index, erd_description = next(
        (i, p) for i, p in enumerate(doc.paragraphs)
        if p.text.strip().startswith("Bảng Songs liên kết với Genres, Tags, Comments")
    )
    replace_paragraph(
        erd_description,
        "Bảng Songs liên kết với Genres, Tags, Song_Comments, Favorites, Playlists và Albums. Bảng Song_Listen_History lưu sự kiện nghe nhạc để phục vụ tính toán xu hướng. Bảng music_analysis_history lưu thông tin phân tích nhạc tham khảo của người dùng và có thể liên kết với Genres. Nhóm thanh toán gồm Packages, Orders, Transactions và Payment_Logs, giúp quản lý gói token, trạng thái đơn hàng và lịch sử giao dịch. Các bảng trung gian Song_Tags, Playlist_Songs và Album_Songs hỗ trợ các quan hệ nhiều - nhiều trong hệ thống.",
    )
    replace_paragraph(
        find_paragraph(doc, "Giao diện tạo nhạc AI cho phép người dùng nhập prompt"),
        "Giao diện tạo nhạc AI được tổ chức theo Wizard gồm các bước nhu cầu, bối cảnh, âm thanh, giọng hát và lời nhạc, sau đó xác nhận yêu cầu. Người dùng lựa chọn provider AI đang khả dụng, thiết lập thông tin bài hát và gửi yêu cầu tạo nhạc. Giao diện hiển thị số token, trạng thái provider, trạng thái xử lý và danh sách bài hát đã tạo. Người dùng có thể theo dõi các trạng thái PENDING, COMPLETED, FAILED hoặc CANCELLED. Giao diện cũng hỗ trợ tải tệp nhạc tham khảo để nhận diện thể loại và áp dụng thể loại gợi ý vào yêu cầu tạo nhạc khi dịch vụ phân tích được cấu hình.",
    )
    replace_paragraph(
        find_paragraph(doc, "3.2.6. Sơ đồ hoạt động mua token qua SePay hoặc VNPay"),
        "Sơ đồ hoạt động mua token qua SePay hoặc VNPay",
    )
    replace_paragraph(
        find_paragraph(doc, "Thông qua cơ chế WebSocket, tin nhắn và trạng thái trực tuyến"),
        "Thông qua cơ chế WebSocket, tin nhắn và trạng thái trực tuyến được truyền đến người nhận gần như ngay lập tức. Nếu người nhận không trực tuyến, hệ thống vẫn lưu tin nhắn và tạo thông báo để người nhận có thể xem sau.",
    )
    erd_image_paragraph = next(
        p for i, p in enumerate(doc.paragraphs)
        if p.text == "" and len(p._p.xpath(".//w:drawing")) > 0 and i > erd_description_index
    )
    insert_figure_caption_after(
        erd_image_paragraph,
        "Sơ đồ thực thể liên kết (ERD) của hệ thống WebMusicAI",
        15,
    )

    # Bảng 3.2: đặc tả Use Case tạo nhạc AI.
    use_case_ai = doc.tables[1]
    field_rows = {row.cells[0].text.strip(): row.cells[1] for row in use_case_ai.rows if len(row.cells) > 1}
    set_cell(field_rows["Tác nhân"], "Người dùng, Dịch vụ sinh nhạc AI")
    set_cell(field_rows["Mục tiêu"], "Cho phép người dùng tạo bài hát từ yêu cầu đã thiết lập, lựa chọn provider AI đang khả dụng và lưu kết quả vào hệ thống.")
    set_cell(field_rows["Điều kiện tiên quyết"], "Người dùng đã đăng nhập, có ít nhất 01 token và provider AI được chọn đang sẵn sàng.")
    set_cell(field_rows["Luồng sự kiện chính"], "1. Người dùng mở giao diện tạo nhạc và lựa chọn provider AI.\n2. Người dùng thiết lập nhu cầu, bối cảnh, thể loại, nhạc cụ, giọng hát, lời nhạc, tiêu đề và ghi chú bổ sung.\n3. Hệ thống kiểm tra dữ liệu đầu vào, trạng thái provider và số dư token.\n4. Hệ thống trừ 01 token, ghi Transaction và tạo Song ở trạng thái PENDING.\n5. Hệ thống đưa yêu cầu vào tác vụ xử lý nền.\n6. MusicGeneratorService chọn provider phù hợp thông qua MusicProviderRegistry và gửi yêu cầu sinh nhạc.\n7. Khi nhận kết quả thành công, hệ thống lưu tệp âm thanh, cập nhật thông tin bài hát và chuyển trạng thái sang COMPLETED.\n8. Hệ thống trả trạng thái mới nhất về giao diện; bài hát công khai có thể tạo thông báo cho người theo dõi tác giả.")
    set_cell(field_rows["Luồng thay thế/ngoại lệ"], "Provider AI không sẵn sàng: hệ thống từ chối yêu cầu và không trừ token. Không đủ token: hệ thống từ chối yêu cầu tạo nhạc. Provider AI xử lý lỗi: hệ thống cập nhật bài hát thành FAILED và hoàn lại 01 token. Người dùng hủy yêu cầu khi bài hát còn PENDING: hệ thống cập nhật CANCELLED và hoàn token.")
    set_cell(field_rows["Hậu điều kiện"], "Bài hát hoàn thành được lưu vào thư viện người dùng; thông tin provider, trạng thái bài hát và lịch sử biến động token được ghi nhận.")

    # Bảng 3.8: mô tả đúng các bảng được bổ sung trong main mới.
    song_data_table = next(t for t in doc.tables if any(row.cells[0].text.strip() == "Songs" for row in t.rows))
    for row in song_data_table.rows:
        if row.cells[0].text.strip() == "Songs":
            set_cell(row.cells[1], "Lưu bài hát được tạo: tiêu đề, prompt, đường dẫn tệp audio, lời bài hát, trạng thái, quyền công khai, ảnh bìa, số lượt nghe, thông tin remix, provider sinh nhạc, thông tin tác vụ và người tạo.")
            break
    append_row(song_data_table, ["Song_Listen_History", "Lưu sự kiện nghe bài hát theo thời gian; phục vụ tính toán xu hướng bài hát, có thể gắn với người nghe khi đã đăng nhập."])
    append_row(song_data_table, ["music_analysis_history", "Lưu thông tin phân tích nhạc tham khảo: người dùng, tên tệp, mã băm, thể loại nhận diện, độ tin cậy, thể loại khớp trong hệ thống và thời điểm phân tích; không lưu tệp gốc."])

    # Chương 4
    replace_paragraph(
        find_paragraph(doc, "Hệ thống tích hợp các dịch vụ bên ngoài gồm"),
        "Hệ thống tích hợp các dịch vụ bên ngoài gồm dịch vụ sinh nhạc AI, dịch vụ phân tích thể loại nhạc, Google OAuth2, SMTP gửi email, SePay và VNPay Sandbox. Dịch vụ sinh nhạc được tổ chức theo cơ chế provider, có thể tích hợp AudioCraft, ACE-Step, MusicAPI.ai hoặc Suno API tùy cấu hình. Các chức năng thời gian thực như chat, thông báo và trạng thái trực tuyến được hỗ trợ thông qua WebSocket.",
    )
    creation_anchor = find_paragraph(doc, "Chức năng tạo nhạc AI được triển khai chủ yếu")
    creation_p2 = find_paragraph(doc, "Khi người dùng gửi yêu cầu tạo nhạc, hệ thống kiểm tra")
    creation_p3 = find_paragraph(doc, "MusicJobService thực hiện xử lý bất đồng bộ")
    creation_p4 = find_paragraph(doc, "Người dùng cũng có thể remix bài hát có sẵn")
    replace_paragraph(
        creation_anchor,
        "Chức năng tạo nhạc AI được triển khai chủ yếu trong SongRestController, SongGenerationService, MusicJobService, MusicGeneratorService, MusicProviderRegistry, các lớp provider sinh nhạc và AudioStorageService. Người dùng thực hiện tạo nhạc thông qua giao diện Wizard gồm các bước nhu cầu, bối cảnh, âm thanh, giọng hát và lời nhạc, sau đó xác nhận yêu cầu.",
    )
    replace_paragraph(
        creation_p2,
        "Khi người dùng gửi yêu cầu, hệ thống kiểm tra trạng thái đăng nhập, số token, nội dung đầu vào và trạng thái sẵn sàng của provider AI được chọn. Nếu hợp lệ, SongGenerationService tạo bản ghi Song ở trạng thái PENDING, trừ token theo quy tắc nghiệp vụ và ghi nhận Transaction.",
    )
    replace_paragraph(
        creation_p3,
        "MusicJobService thực hiện xử lý bất đồng bộ nhằm tránh làm gián đoạn giao diện người dùng. MusicGeneratorService sử dụng MusicProviderRegistry để chọn provider theo yêu cầu. Các provider hiện được tổ chức dưới dạng các lớp riêng như AudioCraftMusicProvider, AceStepMusicProvider, MusicApiMusicProvider và SunoMusicProvider. Khi provider trả về dữ liệu âm thanh, AudioStorageService lưu tệp theo content type nhận được và hệ thống cập nhật bài hát sang trạng thái COMPLETED. Nếu quá trình xử lý thất bại, bài hát được chuyển sang FAILED và token được hoàn lại theo quy tắc nghiệp vụ.",
    )
    replace_paragraph(
        creation_p4,
        "Người dùng cũng có thể tạo phiên bản remix từ bài hát đã hoàn thành mà họ sở hữu hoặc bài hát công khai. Hệ thống tạo bài hát mới, lưu quan hệ với bài hát gốc thông qua trường parent_id và xử lý yêu cầu remix thông qua provider AI được chọn.",
    )
    replace_paragraph(
        find_paragraph(doc, "ReportRestController cung cấp dữ liệu phục vụ dashboard"),
        "ReportRestController cung cấp dữ liệu phục vụ dashboard như tổng số người dùng, số bài hát, doanh thu, số lượng đơn hàng theo trạng thái, bài hát được yêu thích, lịch sử nghe và xu hướng hoạt động. AdminRestController hỗ trợ tìm kiếm, lọc dữ liệu và phân trang; lịch sử giao dịch và Payment Log có thể được lọc theo từ khóa, loại giao dịch và khoảng thời gian. Các số liệu này giúp quản trị viên theo dõi tình hình hoạt động và hỗ trợ việc đối soát giao dịch.",
    )
    replace_paragraph(
        find_paragraph(doc, "Hệ thống sử dụng DatabaseCleanupTask để thực hiện"),
        "Hệ thống sử dụng DatabaseCleanupTask để thực hiện các tác vụ dọn dẹp dữ liệu tự động. Cứ mỗi 10 phút, tác vụ kiểm tra các bài hát ở trạng thái PENDING quá 30 phút; những bài hát này được cập nhật sang FAILED và hoàn token cho người dùng nếu phù hợp. Đồng thời, các đơn thanh toán ở trạng thái PENDING quá 15 phút được chuyển sang EXPIRED.",
    )
    replace_paragraph(
        find_paragraph(doc, "Bên cạnh đó, hệ thống xóa các bài hát ở trạng thái FAILED"),
        "Bên cạnh đó, hệ thống xóa các bài hát ở trạng thái FAILED quá 07 ngày vào 02:00 hằng ngày. Các tác vụ này giúp giảm dữ liệu không cần thiết và bảo đảm trạng thái bài hát, token và đơn hàng được cập nhật ổn định.",
    )

    # Chương 5
    # The testing-environment table has the shared header "Thành phần" and a Backend row.
    environment = next(
        t for t in doc.tables
        if t.rows and t.rows[0].cells[0].text.strip() == "Thành phần"
        and any(row.cells[0].text.strip() == "Backend" for row in t.rows)
    )
    for row in environment.rows:
        if row.cells[0].text.strip() == "Dịch vụ AI":
            set_cell(row.cells[1], "AudioCraft/ACE-Step trên Colab hoặc API AI ngoài được cấu hình")
            set_cell(row.cells[2], "Kiểm thử trạng thái provider, tạo nhạc bất đồng bộ và xử lý kết quả âm thanh")
            break
    else:
        raise ValueError("Không tìm thấy hàng Dịch vụ AI trong Bảng 5.1")

    test_catalog = next(t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Mã" and t.rows[0].cells[1].text.strip() == "Nhóm chức năng")
    append_row(test_catalog, ["TC-10", "Provider AI", "Kiểm tra provider online/offline", "Chỉ tạo tác vụ và trừ token khi provider sẵn sàng"])
    append_row(test_catalog, ["TC-11", "Phân tích nhạc tham khảo", "Nhận diện thể loại tệp nhạc", "Kiểm tra định dạng, kích thước, kết quả nhận diện và lịch sử phân tích"])

    ai_test_table = next(t for t in doc.tables if t.rows and t.rows[0].cells[0].text.strip() == "Mã" and any("Gửi prompt khi còn token" in c.text for row in t.rows for c in row.cells))
    append_row(ai_test_table, ["TC-03.6", "Chọn provider không sẵn sàng", "Từ chối yêu cầu, không tạo Song PENDING và không trừ token", "Chưa kiểm thử"])
    append_row(ai_test_table, ["TC-03.7", "Tải tệp nhạc tham khảo hợp lệ", "Nhận diện thể loại, lưu lịch sử phân tích và cho phép áp dụng thể loại gợi ý", "Chưa kiểm thử"])

    replace_paragraph(
        find_paragraph(doc, "Các luồng trọng yếu đã được kiểm thử"),
        "Các luồng trọng yếu đã được kiểm thử trên môi trường phát triển và Sandbox. Hệ thống xử lý trạng thái PENDING, COMPLETED, FAILED và CANCELLED; hoàn token đúng luồng khi provider AI thất bại hoặc người dùng hủy tác vụ. SePay và VNPay được tách luồng xác nhận rõ ràng; PaymentService hạn chế cộng token trùng lặp. Với các provider AI hoặc dịch vụ phân tích chưa được cấu hình ở môi trường kiểm thử, kết quả cần được ghi nhận là Chưa kiểm thử thay vì Pass.",
    )
    replace_paragraph(
        find_paragraph(doc, "Chất lượng và thời gian tạo nhạc phụ thuộc vào AI Music API."),
        "Chất lượng đầu ra và thời gian tạo nhạc phụ thuộc vào provider AI đang được cấu hình. Một số provider chạy thông qua Google Colab hoặc API bên ngoài nên có thể không sẵn sàng liên tục, bị giới hạn tài nguyên hoặc thay đổi thời gian phản hồi. Giao dịch REVIEW cần đối soát thủ công. Kiểm thử cần tiếp tục mở rộng theo hướng Unit Test, Integration Test tự động, kiểm thử tải AI/WebSocket và giám sát tập trung khi triển khai thực tế.",
    )
    replace_paragraph(
        find_paragraph(doc, "Chương 5 đã trình bày môi trường, kịch bản"),
        "Chương 5 đã trình bày môi trường, kịch bản và kết quả kiểm thử của các nhóm chức năng trọng yếu. Các test case cần được cập nhật kết quả thực thi, ngày kiểm thử và minh chứng tương ứng trước khi nộp báo cáo chính thức.",
    )

    # Chương 6
    replace_paragraph(
        find_paragraph(doc, "Nhóm đã xây dựng hệ thống WebMusicAI với các nhóm chức năng chính gồm:"),
        "Nhóm đã xây dựng hệ thống WebMusicAI với các nhóm chức năng chính gồm: đăng ký, đăng nhập bằng JWT và Google OAuth2; xác thực OTP để đặt lại mật khẩu; tạo nhạc từ yêu cầu người dùng; lựa chọn provider sinh nhạc AI đang khả dụng; phân tích nhạc tham khảo để gợi ý thể loại; quản lý bài hát và thư viện cá nhân; tương tác cộng đồng; nhắn tin thời gian thực; thanh toán token qua SePay và VNPay Sandbox; cùng chức năng quản trị và thống kê trên Dashboard.",
    )
    replace_paragraph(
        find_paragraph(doc, "WebMusicAI xây dựng quy trình tương đối hoàn chỉnh"),
        "WebMusicAI xây dựng quy trình tương đối hoàn chỉnh từ khâu thiết lập yêu cầu sáng tạo, lựa chọn provider AI, tạo nhạc, lưu trữ, quản lý đến chia sẻ bài hát. Việc xử lý tạo nhạc theo cơ chế bất đồng bộ giúp người dùng vẫn có thể tiếp tục thao tác trên giao diện trong khi hệ thống chờ kết quả từ dịch vụ AI.",
    )
    replace_paragraph(
        find_paragraph(doc, "Chất lượng đầu ra và thời gian xử lý tạo nhạc phụ thuộc vào dịch vụ AI bên ngoài."),
        "Chất lượng đầu ra và thời gian xử lý tạo nhạc phụ thuộc vào provider AI đang được cấu hình. Hệ thống hiện chưa có hạ tầng GPU do nhóm tự vận hành ổn định trong môi trường production; việc tạo nhạc vẫn phụ thuộc vào worker Google Colab hoặc API AI bên ngoài. Vì vậy, khả năng sẵn sàng, tốc độ phản hồi, giới hạn tài nguyên và chi phí có thể thay đổi theo từng provider.",
    )
    replace_paragraph(
        find_paragraph(doc, "Trong thời gian tới, nhóm định hướng triển khai dịch vụ AI"),
        "Trong thời gian tới, nhóm định hướng chuyển các worker AI thử nghiệm trên Colab sang hạ tầng cloud GPU được quản lý, bổ sung hàng đợi tác vụ và cơ chế giám sát quá trình tạo nhạc. Hệ thống cũng có thể mở rộng thêm các phương thức thanh toán, nhà cung cấp OAuth, gói dịch vụ và chức năng đề xuất bài hát cá nhân hóa.",
    )

    # Phụ lục đối chiếu mã nguồn.
    appendix_anchor = find_paragraph(doc, "Sinh nhạc AI: SongRestController tiếp nhận yêu cầu")
    replace_paragraph(
        appendix_anchor,
        "Sinh nhạc AI: SongRestController tiếp nhận yêu cầu tạo nhạc và remix. SongGenerationService kiểm tra token, tạo bài hát ở trạng thái PENDING và trừ token theo quy tắc hệ thống. MusicJobService xử lý tác vụ nền; MusicGeneratorService sử dụng MusicProviderRegistry để chọn provider sinh nhạc phù hợp. Hệ thống hiện tổ chức các provider như AudioCraftMusicProvider, AceStepMusicProvider, MusicApiMusicProvider và SunoMusicProvider. Khi nhận kết quả, AudioStorageService lưu trữ tệp âm thanh và hệ thống cập nhật trạng thái COMPLETED hoặc FAILED. Token được hoàn lại khi tác vụ thất bại, bị hủy hoặc bị treo quá thời gian quy định.",
    )
    insert_after(
        appendix_anchor,
        "Phân tích nhạc tham khảo: MusicAnalysisController tiếp nhận tệp âm thanh tham khảo từ người dùng. MusicReferenceAnalysisService kiểm tra định dạng, kích thước tệp, tạo mã băm để tái sử dụng kết quả đã có và gọi dịch vụ phân tích thể loại nhạc bên ngoài. Kết quả gồm thể loại nhận diện, độ tin cậy và thể loại tương ứng trong hệ thống; lịch sử được lưu tại bảng music_analysis_history, không lưu tệp âm thanh tham khảo gốc.",
        appendix_anchor,
    )
    replace_paragraph(
        find_paragraph(doc, "Quản trị và báo cáo: AdminRestController quản lý"),
        "Quản trị và báo cáo: AdminRestController quản lý người dùng, bài hát, gói token, thể loại, thẻ, đơn hàng, lịch sử thanh toán và duyệt các giao dịch REVIEW. Hệ thống hỗ trợ tìm kiếm, phân trang và lọc dữ liệu; lịch sử giao dịch và Payment Log có thể lọc theo từ khóa, loại giao dịch và khoảng ngày. ReportRestController cung cấp số liệu tổng quan, doanh thu, bài hát được yêu thích, lịch sử nghe, người dùng và dữ liệu tăng trưởng phục vụ dashboard quản trị.",
    )
    replace_paragraph(
        find_paragraph(doc, "Tác vụ nền và thời gian thực: DatabaseCleanupTask"),
        "Tác vụ nền và thời gian thực: DatabaseCleanupTask chạy mỗi 10 phút để cập nhật FAILED và hoàn token cho bài hát PENDING quá 30 phút; đồng thời chuyển đơn PENDING quá 15 phút sang EXPIRED. Tác vụ xóa bài hát FAILED quá 07 ngày chạy lúc 02:00 hằng ngày. WebSocketConfig, ChatWebSocketController và ChatMessageService hỗ trợ gửi tin nhắn, trạng thái trực tuyến và thông báo thời gian thực.",
    )

    replace_paragraph(find_paragraph(doc, "BẢNG PHÂN CÔNG CÔNG VIỆC CỦA NHÓM"), "PHỤ LỤC A. BẢNG PHÂN CÔNG CÔNG VIỆC CỦA NHÓM")
    replace_paragraph(find_paragraph(doc, "PHỤ LỤC : ĐỐI CHIẾU CHỨC NĂNG VỚI MÃ NGUỒN"), "PHỤ LỤC B. ĐỐI CHIẾU CHỨC NĂNG VỚI MÃ NGUỒN")

    set_update_fields_on_open(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
